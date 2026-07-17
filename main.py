"""
doc-formatter Web 平台 - FastAPI 后端
"""
import os
import uuid
import shutil
from datetime import datetime
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi import Request
from fastapi.templating import Jinja2Templates

from engine import process_document, save_to_knowledge_base
from docx import Document
from docx.shared import Pt


def _md_to_docx(md_path):
    """将 Markdown 规范/标杆文档转为临时 .docx，供排版引擎读取"""
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
            level = len(line.split()[0]) - 1
            doc.add_heading(line.lstrip('#').strip(), level=level)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)

    tmp_path = md_path.with_suffix('.docx.tmp')
    doc.save(str(tmp_path))
    return tmp_path


def _clean_tmp(base_path):
    """清理由 _md_to_docx / _pdf_text_to_file 生成的临时文件"""
    for suffix in ('.docx.tmp', '.md.tmp'):
        tmp_path = Path(str(base_path) + suffix) if not str(base_path).endswith(suffix) else Path(str(base_path))
        if tmp_path.exists():
            tmp_path.unlink()


def _pdf_text_to_file(pdf_path):
    """从 PDF 提取文字并写出到临时 Markdown 文件，供规范/标杆模式使用"""
    try:
        import pymupdf
    except ImportError:
        return None
    doc = pymupdf.open(str(pdf_path))
    text = ''
    for page in doc:
        text += page.get_text()
    doc.close()

    tmp_path = pdf_path.with_suffix('.md.tmp')
    with open(str(tmp_path), 'w', encoding='utf-8') as f:
        f.write(text)
    return tmp_path

app = FastAPI(title="文档自动排版平台")

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / 'uploads'
STATIC_DIR = BASE_DIR / 'static'
TEMPLATE_DIR = BASE_DIR / 'templates'

# 确保目录存在
for d in [UPLOAD_DIR / 'target', UPLOAD_DIR / 'specification',
          UPLOAD_DIR / 'benchmark', UPLOAD_DIR / 'output']:
    d.mkdir(parents=True, exist_ok=True)

app.mount('/static', StaticFiles(directory=str(STATIC_DIR)), name='static')
templates = Jinja2Templates(directory=str(TEMPLATE_DIR))


@app.get('/', response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse('index.html', {'request': request})


@app.post('/api/upload/{file_type}')
async def upload_file(file_type: str, file: UploadFile = File(...)):
    """上传文件"""
    valid_types = {'target', 'specification', 'benchmark'}

    if file_type not in valid_types:
        raise HTTPException(400, f'无效的文件类型，支持: {valid_types}')

    # target 只允许 .docx；规范/标杆允许 .docx/.md/.pdf
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if file_type == 'target' and ext != 'docx':
        raise HTTPException(400, '待排版文档仅支持 .docx 格式')
    if file_type in ('specification', 'benchmark') and ext not in ('docx', 'md', 'txt', 'pdf'):
        raise HTTPException(400, '规范/标杆文档支持 .docx / .md / .pdf 格式')

    ext = file.filename.rsplit('.', 1)[-1]
    unique_name = f'{uuid.uuid4().hex}_{datetime.now().strftime("%Y%m%d%H%M%S")}.{ext}'
    dest = UPLOAD_DIR / file_type / unique_name

    with open(dest, 'wb') as f:
        content = await file.read()
        f.write(content)

    return {
        'filename': file.filename,
        'stored_name': unique_name,
        'path': str(dest),
        'size': len(content),
    }


@app.post('/api/process')
async def handle_process(data: dict):
    """
    处理文档排版
    body: {
        "mode": "instruction" | "specification" | "benchmark",
        "target_file": "上传后的文件名",
        "instructions": "格式要求（模式一用）",
        "spec_file": "规范文件名（模式二用）",
        "benchmark_file": "标杆文档名（模式三用）",
        "save_to_kb": true/false  # 是否保存到知识库
    }
    """
    mode = data.get('mode')
    target_file = data.get('target_file')
    instructions = data.get('instructions', '')
    spec_file = data.get('spec_file')
    benchmark_file = data.get('benchmark_file')
    save_to_kb = data.get('save_to_kb', False)
    original_filename = data.get('original_filename', target_file)  # 用户原始文件名

    if not target_file:
        raise HTTPException(400, '请上传待排版的文档')

    target_path = UPLOAD_DIR / 'target' / target_file
    if not target_path.exists():
        raise HTTPException(404, '目标文件未找到，请重新上传')

    # 生成输出路径（用原始文件名）
    base_name = os.path.splitext(original_filename)[0]
    output_name = f'{base_name}-排版后.docx'
    output_path = UPLOAD_DIR / 'output' / output_name

    try:
        # 确定规范/标杆文件路径（供读取文本内容）
        spec_path = None
        source_name = ''

        if mode in ('specification', 'benchmark'):
            file_key = 'spec_file' if mode == 'specification' else 'benchmark_file'
            uploaded_file = data.get(file_key)

            if uploaded_file:
                # 上传目录
                sub_dir = 'specification' if mode == 'specification' else 'benchmark'
                candidate = UPLOAD_DIR / sub_dir / uploaded_file
                if candidate.exists():
                    spec_path = candidate
                    source_name = f'用户上传 ({uploaded_file})'
                else:
                    # 知识库
                    kb_sub = '01-格式规范' if mode == 'specification' else '02-标杆文档'
                    kb_candidate = BASE_DIR / 'knowledge_base' / kb_sub / uploaded_file
                    if kb_candidate.exists():
                        spec_path = kb_candidate
                        source_name = f'知识库 ({uploaded_file})'
                    else:
                        # 也查原始知识库（从配置读取）
                        try:
                            import yaml
                            with open(BASE_DIR / 'config.yaml', 'r', encoding='utf-8') as _cf:
                                _cfg = yaml.safe_load(_cf)
                            _kb_root = Path(_cfg.get('knowledge_base', {}).get('path', 'knowledge_base'))
                            if not _kb_root.is_absolute():
                                _kb_root = BASE_DIR / _kb_root
                            alt = _kb_root / kb_sub / uploaded_file
                            if alt.exists():
                                spec_path = alt
                                source_name = f'知识库 ({uploaded_file})'
                        except Exception:
                            pass

            # 未上传时用默认
            if spec_path is None:
                if mode == 'specification':
                    kb_spec = _find_default_spec()
                else:
                    kb_spec = _find_default_benchmark()
                if kb_spec:
                    spec_path = kb_spec
                    source_name = kb_spec.name
                else:
                    raise HTTPException(400, f'未找到默认{"规范" if mode=="specification" else "标杆"}文件')

        # 调用统一引擎
        result = process_document(
            doc_path=str(target_path),
            output_path=str(output_path),
            mode=mode,
            spec_path=str(spec_path) if spec_path else None,
            user_instructions=instructions,
            original_filename=original_filename,
        )

        # 处理失败
        if not result.get('success'):
            raise HTTPException(500, result.get('error', '排版处理失败'))

        # 补全来源信息
        result['report']['source'] = source_name or f'LLM-{mode}'
        result['output_file'] = output_name

        # 保存到知识库
        if save_to_kb and spec_path:
            kb_type = 'specification' if mode == 'specification' else 'benchmark'
            kb_target = save_to_knowledge_base(str(spec_path), kb_type)
            if kb_target:
                kb_sub = '01-格式规范' if mode == 'specification' else '02-标杆文档'
                result['report']['knowledge_base_update'] = f'已保存至知识库 {kb_sub}/{os.path.basename(kb_target)}'

        # 自动保存（从配置读取路径）
        try:
            import yaml
            with open(BASE_DIR / 'config.yaml', 'r', encoding='utf-8') as _cf:
                _cfg = yaml.safe_load(_cf)
            _save_root = Path(_cfg.get('auto_save', {}).get('path', 'uploads/output'))
            if not _save_root.is_absolute():
                _save_root = BASE_DIR / _save_root
            _save_root.mkdir(parents=True, exist_ok=True)
            src = UPLOAD_DIR / 'output' / output_name
            if src.exists():
                import shutil
                dst = _save_root / output_name
                shutil.copy2(str(src), str(dst))
                result['auto_saved_to'] = str(dst)
        except Exception:
            pass

        return result

    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        print(f"!!! 排版错误: {tb}")
        raise HTTPException(500, f'排版处理出错: {str(e)}\n{tb[:500]}')


@app.get('/api/download/{filename}')
async def download(filename: str):
    """下载排版后的文档"""
    file_path = UPLOAD_DIR / 'output' / filename
    if not file_path.exists():
        raise HTTPException(404, '文件未找到')
    return FileResponse(
        str(file_path),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=filename,
    )


@app.get('/api/knowledge-base')
async def list_knowledge_base():
    """列出知识库中的格式规范和标杆文档"""
    kb_path = BASE_DIR / 'knowledge_base'
    result = {'specifications': [], 'benchmarks': []}

    spec_dir = kb_path / '01-格式规范'
    if spec_dir.exists():
        for f in sorted(spec_dir.iterdir()):
            if f.suffix in ('.md', '.docx', '.txt', '.pdf'):
                ext_info = '📋' if f.suffix in ('.md', '.docx', '.txt') else '📕'
                result['specifications'].append({
                    'name': f.name,
                    'ext': f.suffix,
                    'icon': ext_info,
                    'modified': datetime.fromtimestamp(f.stat().st_mtime).strftime('%Y-%m-%d %H:%M'),
                })

    bm_dir = kb_path / '02-标杆文档'
    if bm_dir.exists():
        for f in sorted(bm_dir.iterdir()):
            if f.suffix in ('.md', '.docx', '.txt', '.pdf'):
                ext_info = '🏷️' if f.suffix in ('.md', '.docx', '.txt') else '📕'
                result['benchmarks'].append({
                    'name': f.name,
                    'ext': f.suffix,
                    'icon': ext_info,
                    'modified': datetime.fromtimestamp(f.stat().st_mtime).strftime('%Y-%m-%d %H:%M'),
                })

    return result


@app.get('/api/debug')
async def debug():
    """调试端点：检查环境变量配置"""
    import os
    result = {
        'DEEPSEEK_API_KEY': bool(os.environ.get('DEEPSEEK_API_KEY')),
        'DEEPSEEK_KEY': bool(os.environ.get('DEEPSEEK_KEY')),
        'DEEPSEEK_APIKEY': bool(os.environ.get('DEEPSEEK_APIKEY')),
        'has_config_file': os.path.exists(os.path.join(os.path.dirname(__file__), 'config.yaml')),
    }
    # 读取当前实际加载的 Key（只显示前5位）
    try:
        import yaml
        with open(os.path.join(os.path.dirname(__file__), 'config.yaml')) as f:
            cfg = yaml.safe_load(f)
        key = cfg.get('llm', {}).get('api_key', '')
        result['config_file_key_preview'] = key[:5] + '...' if key else '(empty)'
    except Exception as e:
        result['config_read_error'] = str(e)
    return result


@app.get('/api/knowledge-base/download/{category}/{filename}')
async def download_kb_file(category: str, filename: str):
    """下载知识库中的文件"""
    category_map = {'spec': '01-格式规范', 'benchmark': '02-标杆文档'}
    dir_name = category_map.get(category)
    if not dir_name:
        raise HTTPException(400, '无效的分类')
    file_path = BASE_DIR / 'knowledge_base' / dir_name / filename
    if not file_path.exists():
        # 也检查原始知识库目录
        alt_path = Path('C:/Users/王逸朴/Desktop/AI agent competition/知识库') / dir_name / filename
        if alt_path.exists():
            file_path = alt_path
    if not file_path.exists():
        raise HTTPException(404, '文件未找到')
    return FileResponse(
        str(file_path),
        filename=filename,
    )


def _find_default_spec():
    """查找知识库中默认的格式规范（优先用 .docx 样式库文件，其次 .md）"""
    spec_dir = BASE_DIR / 'knowledge_base' / '01-格式规范'
    if spec_dir.exists():
        files = list(spec_dir.glob('*.*'))
        docx_files = [f for f in files if f.suffix == '.docx']
        if docx_files:
            return docx_files[0]
        md_files = [f for f in files if f.suffix == '.md']
        if md_files:
            return md_files[0]
        pdf_files = [f for f in files if f.suffix == '.pdf']
        if pdf_files:
            return pdf_files[0]
        if files:
            return files[0]
    return None


def _find_default_benchmark():
    """查找知识库中默认的标杆文档（优先用 .docx，其次 .md）"""
    bm_dir = BASE_DIR / 'knowledge_base' / '02-标杆文档'
    if bm_dir.exists():
        files = list(bm_dir.glob('*.*'))
        docx_files = [f for f in files if f.suffix == '.docx']
        if docx_files:
            return docx_files[0]
        md_files = [f for f in files if f.suffix == '.md']
        if md_files:
            return md_files[0]
        pdf_files = [f for f in files if f.suffix == '.pdf']
        if pdf_files:
            return pdf_files[0]
        if files:
            return files[0]
    return None


if __name__ == '__main__':
    import uvicorn
    port = int(os.environ.get('PORT', 8000))
    uvicorn.run(app, host='0.0.0.0', port=port)
