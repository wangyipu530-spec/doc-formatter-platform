"""Zeabur 端到端测试"""
import urllib.request, json, tempfile, os, sys
from docx import Document
from docx.shared import Pt

doc = Document()
h = doc.add_paragraph(); r = h.add_run('测试标题'); r.font.size = Pt(26)
doc.add_paragraph('正文内容第一行。')
doc.add_paragraph('正文内容第二行。')

p = os.path.join(tempfile.gettempdir(), 'zt2.docx')
doc.save(p)

with open(p, 'rb') as f: data = f.read()
b = '----B'
body = (f'--{b}\r\nContent-Disposition: form-data; name="file"; filename="t.docx"\r\n'
        'Content-Type: application/octet-stream\r\n\r\n').encode() + data + f'\r\n--{b}--\r\n'.encode()
req = urllib.request.Request(
    'https://docformatter.preview.aliyun-zeabur.cn/api/upload/target',
    data=body, headers={'Content-Type': f'multipart/form-data; boundary={b}'})
r = json.loads(urllib.request.urlopen(req, timeout=15).read())
print(f'上传: {r["stored_name"][:20]}...')

# 排版
payload = json.dumps({
    'mode': 'instruction', 'target_file': r['stored_name'],
    'original_filename': 't.docx',
    'instructions': '标题改为黑体三号加粗居中，正文宋体小四1.5倍行距首行缩进2字符',
}).encode()
req2 = urllib.request.Request(
    'https://docformatter.preview.aliyun-zeabur.cn/api/process',
    data=payload, headers={'Content-Type': 'application/json'})
try:
    r2 = json.loads(urllib.request.urlopen(req2, timeout=120).read())
    if r2.get('success'):
        print('✅ 排版成功')
        for c in r2['report']['changes']:
            print(f'  {c}')
        print(f'文件: {r2["output_file"]}')
    else:
        print(f'❌ 失败: {r2.get("error", "")[:200]}')
except urllib.error.HTTPError as e:
    body = e.read().decode()
    print(f'❌ HTTP {e.code}: {body[:500]}')
except Exception as e:
    print(f'❌ 错误: {e}')
finally:
    os.remove(p)
