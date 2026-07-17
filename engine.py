"""
排版引擎 — 执行层
不再负责「理解」格式要求（交给 LLM），只负责两件事：
  1. 提取文档格式快照（供 LLM 分析）
  2. 执行 LLM 返回的 JSON 指令（操作 .docx XML）
"""
import os
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from llm_adapter import extract_document_snapshot, analyze_and_plan


# ─── 字体映射 ───
FONT_MAP = {
    '宋体': 'SimSun', '仿宋': 'FangSong', '黑体': 'SimHei',
    '楷体': 'KaiTi', '微软雅黑': 'Microsoft YaHei', 'Arial': 'Arial',
    'Times New Roman': 'Times New Roman',
    '小标宋': 'SimSun', '小标宋体': 'SimSun',
}

SIZE_MAP = {
    '初号': 42, '小初': 36, '一号': 26, '小一': 24,
    '二号': 22, '小二': 18, '三号': 16, '小三': 15,
    '四号': 14, '小四': 12, '五号': 10.5, '小五': 9,
    '六号': 7.5, '小六': 6.5,
}


def _set_cn_font(run, font_name):
    """同时设置西文和中文字体"""
    try:
        run.font.name = font_name
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), font_name)
    except:
        pass


def _is_heading(para):
    name = para.style.name.lower() if para.style.name else ''
    return name.startswith('heading')


# ═══════════════════════════════════════════════
# 核心函数：执行 LLM 指令
# ═══════════════════════════════════════════════

def _apply_paragraph_instruction(doc, instruction):
    """对指定段落应用一条格式指令"""
    indices = instruction.get('indices', [])
    font_val = instruction.get('font')
    size_val = instruction.get('size_pt')
    bold_val = instruction.get('bold')
    italic_val = instruction.get('italic')
    color_val = instruction.get('color')
    alignment_val = instruction.get('alignment')
    line_spacing_val = instruction.get('line_spacing')
    indent_val = instruction.get('first_line_indent_cm')

    # 如果 indices 为空，遍历所有非空段落
    paras_to_modify = []
    if not indices:
        paras_to_modify = [p for p in doc.paragraphs if p.text.strip()]
    else:
        for idx in indices:
            if 0 <= idx < len(doc.paragraphs):
                paras_to_modify.append(doc.paragraphs[idx])

    count = 0
    for para in paras_to_modify:
        if not para.text.strip():
            continue

        # 对齐
        if alignment_val:
            align_map = {
                'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                'BOTH': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            try:
                para.alignment = align_map.get(alignment_val.upper())
            except:
                pass

        # 行距
        if line_spacing_val is not None:
            try:
                para.paragraph_format.line_spacing = line_spacing_val
            except:
                pass

        # 首行缩进
        if indent_val is not None:
            try:
                para.paragraph_format.first_line_indent = Cm(indent_val)
            except:
                pass

        # 每个 run 的字体属性
        for run in para.runs:
            if font_val:
                _set_cn_font(run, font_val)
            if size_val is not None:
                try:
                    run.font.size = Pt(size_val)
                except:
                    pass
            if bold_val is not None:
                try:
                    run.font.bold = bold_val
                except:
                    pass
            if italic_val is not None:
                try:
                    run.font.italic = italic_val
                except:
                    pass
            if color_val:
                try:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(*bytes.fromhex(color_val.lstrip('#')))
                except:
                    pass

        count += 1

    return count


def _apply_global_instructions(doc, global_inst):
    """应用全局设置（页边距等）"""
    margins = global_inst.get('margins', {})
    if margins:
        for section in doc.sections:
            if 'top_cm' in margins and margins['top_cm'] is not None:
                section.top_margin = Cm(margins['top_cm'])
            if 'bottom_cm' in margins and margins['bottom_cm'] is not None:
                section.bottom_margin = Cm(margins['bottom_cm'])
            if 'left_cm' in margins and margins['left_cm'] is not None:
                section.left_margin = Cm(margins['left_cm'])
            if 'right_cm' in margins and margins['right_cm'] is not None:
                section.right_margin = Cm(margins['right_cm'])
        return True
    return False


def _build_report(mode, changes, issues, suggestions, original_filename, source):
    """生成排版变更报告"""
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    return {
        'success': True,
        'report': {
            'filename': original_filename,
            'mode': mode,
            'source': source,
            'time': now,
            'changes': changes,
            'warnings': issues,
            'suggestions': suggestions,
        }
    }


# ═══════════════════════════════════════════════
# 三种排版模式的统一入口
# ═══════════════════════════════════════════════

def process_document(doc_path, output_path, mode, spec_path=None,
                     user_instructions='', original_filename='document.docx'):
    """
    统一排版入口。

    流程：
    1. 打开文档
    2. 提取格式快照
    3. 读取规范内容（如果有）
    4. 调 LLM 分析差异，返回指令
    5. 执行指令
    6. 保存 + 返回报告
    """
    doc = Document(doc_path)
    changes = []
    issues = []
    suggestions = []

    # 1. 提取格式快照
    snapshot = extract_document_snapshot(doc)

    # 2. 读取规范内容
    spec_text = ''
    if spec_path and os.path.exists(spec_path):
        ext = os.path.splitext(spec_path)[1].lower()
        if ext == '.pdf':
            try:
                import pymupdf
                pdf_doc = pymupdf.open(spec_path)
                spec_text = '\n'.join([page.get_text() for page in pdf_doc])
                pdf_doc.close()
            except ImportError:
                spec_text = f'[PDF 文件，但未安装 pymupdf: {os.path.basename(spec_path)}]'
            except Exception as e:
                spec_text = f'[PDF 读取失败: {str(e)}]'
        elif ext == '.docx':
            try:
                spec_doc = Document(spec_path)
                spec_text = '\n'.join([p.text for p in spec_doc.paragraphs if p.text.strip()])
            except Exception as e:
                spec_text = f'[DOCX 读取失败: {str(e)}]'
        else:
            try:
                with open(spec_path, 'r', encoding='utf-8') as f:
                    spec_text = f.read()
            except:
                try:
                    spec_doc = Document(spec_path)
                    spec_text = '\n'.join([p.text for p in spec_doc.paragraphs if p.text.strip()])
                except:
                    spec_text = f'[无法读取规范文件: {os.path.basename(spec_path)}]'
    elif spec_path:
        spec_text = f'[规范文件不存在: {spec_path}]'

    # 3. 调 LLM 分析
    mode_names = {
        'instruction': '按指令修改',
        'specification': '按规范排版',
        'benchmark': '按标杆文档排版',
    }
    mode_name = mode_names.get(mode, mode)

    try:
        plan = analyze_and_plan(mode, snapshot, spec_text, user_instructions)
    except Exception as e:
        return {
            'success': False,
            'error': f'LLM 分析失败: {str(e)}',
        }

    # 4. 执行段落指令
    para_instructions = plan.get('paragraph_instructions', [])
    for inst in para_instructions:
        count = _apply_paragraph_instruction(doc, inst)
        # 生成易读的变更描述
        desc_parts = []
        if inst.get('font'): desc_parts.append(f"字体={inst['font']}")
        if inst.get('size_pt'): desc_parts.append(f"字号={inst['size_pt']}pt")
        if inst.get('bold') is True: desc_parts.append('加粗')
        if inst.get('bold') is False: desc_parts.append('取消加粗')
        if inst.get('italic') is True: desc_parts.append('斜体')
        if inst.get('alignment'): desc_parts.append(f"对齐={inst['alignment']}")
        if inst.get('line_spacing'): desc_parts.append(f"行距={inst['line_spacing']}倍")
        if inst.get('first_line_indent_cm'): desc_parts.append(f"首行缩进{inst['first_line_indent_cm']}cm")

        idx_desc = f"段落{inst.get('indices', [])}" if inst.get('indices') else '全文'
        if desc_parts:
            changes.append(f"[段落] {idx_desc}（共{count}处）：{'、'.join(desc_parts)}")

    # 5. 执行全局指令
    global_inst = plan.get('global_instructions', {})
    if _apply_global_instructions(doc, global_inst):
        margins = global_inst.get('margins', {})
        if margins:
            parts = []
            for k, label in [('top_cm', '上'), ('bottom_cm', '下'), ('left_cm', '左'), ('right_cm', '右')]:
                v = margins.get(k)
                if v is not None:
                    parts.append(f'{label}{v}cm')
            if parts:
                changes.append(f"[页边距] 已调整为{'、'.join(parts)}")

    # 6. 收集 LLM 的分析信息
    analysis = plan.get('analysis', '')
    if analysis:
        # 把 analysis 作为第一个变更项（概要）
        changes.insert(0, f"[AI 分析] {analysis}")

    issues.extend(plan.get('issues', []))
    suggestions.extend(plan.get('suggestions', []))

    if not changes:
        changes.append("[提示] 文档格式已符合要求，无需修改")

    # 7. 保存
    doc.save(output_path)

    # 构造来源描述
    if mode == 'instruction':
        source = f'LLM分析用户指令: {user_instructions[:60]}'
    elif spec_path:
        source = f'LLM分析规范文件: {os.path.basename(spec_path)}'
    else:
        source = f'LLM分析 ({mode_name})'

    return _build_report(mode_name, changes, issues, suggestions, original_filename, source)


# ─── 知识库管理 ───
KNOWLEDGE_BASE_PATH = os.path.join(os.path.dirname(__file__), 'knowledge_base')


def save_to_knowledge_base(uploaded_file_path, file_type='specification'):
    """将上传的文件保存到知识库"""
    if not os.path.exists(uploaded_file_path):
        return None
    from datetime import datetime
    import shutil

    today = datetime.now().strftime('%Y%m%d')
    filename = os.path.basename(uploaded_file_path)

    if file_type == 'specification':
        target_dir = os.path.join(KNOWLEDGE_BASE_PATH, '01-格式规范')
        new_name = f'用户上传_{today}_{filename}'
    elif file_type == 'benchmark':
        target_dir = os.path.join(KNOWLEDGE_BASE_PATH, '02-标杆文档')
        new_name = f'标杆文档_{today}_{filename}'
    else:
        return None

    os.makedirs(target_dir, exist_ok=True)
    dest_path = os.path.join(target_dir, new_name)

    if os.path.exists(dest_path):
        base, ext = os.path.splitext(new_name)
        counter = 1
        while os.path.exists(dest_path):
            dest_path = os.path.join(target_dir, f'{base}_{counter}{ext}')
            counter += 1

    shutil.copy2(uploaded_file_path, dest_path)
    return dest_path
